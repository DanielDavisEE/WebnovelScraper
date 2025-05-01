"""----- Create Document -----"""
import docx

def create_docx(title, chapter_dict):
    document = docx.Document('StyleDoc.docx')

    # Title Page
    document.add_paragraph('\n' * 5)
    document.add_heading(title, 0)
    document.add_page_break()

    # Contents Page
    document.add_heading('Contents', 1)

    for key, value in chapter_dict.items():
        document.add_paragraph(f"{key}. {value['title']}")

    document.add_page_break()

    # Write all chapters
    for key, value in chapter_dict.items():
        document.add_heading(f'Chapter {key}', level=2)
        # document.add_heading(f"{key}. {value['title']}", level=2)

        for line in value['text']:
            document.add_paragraph(line)

        document.add_page_break()

    document.save(f"{title}\\{title}.docx")
    print('Document successfully created!')
