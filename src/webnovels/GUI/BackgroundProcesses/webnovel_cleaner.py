import docx, time
import re
from bs4 import BeautifulSoup
#from spellchecker import SpellChecker
#from mySpellCheck import *
# –

CHAPTER_LIMIT = 0
TITLE = "I Alone Level Up"
URL_BASE = 'https://wuxiaworld.site/novel/only-i-level-up/'

if ' ' in TITLE.strip(' \n\t'):
    tmp = TITLE[0]
    for i in range(1, len(TITLE)):
        if TITLE[i-1] == ' ':
            tmp = tmp + TITLE[i]

filename = f"{tmp}Dict.txt"

'''
spell = SpellChecker()
spell.word_frequency.load_text_file(filename, encoding=u'utf-8')

dict_infile = open(filename, 'a')

# Dictionary to replace incorrect words
replacementDict = {}
with open(f"{TITLE}\\ReplacementDict.txt", 'r', encoding='utf-8') as file:
    word_list = file.readlines()
    for item in word_list:
        try:
            word, replacement = item.split(maxsplit=1)
        except ValueError as e:
            print(item)
            raise ValueError(e)
        replacementDict[word] = replacement.strip('\n')
        
# Replace words that have been censored using periods ie b.l.o.o.d.y => bloody
# Could be within other words
censoredDict = {}
with open(f"{TITLE}\\CensoredWords.txt", 'r', encoding='utf-8') as file:
    word_list = file.readlines()
    for item in word_list:
        word, replacement = item.split('|', maxsplit=1)
        censoredDict[word] = replacement.strip('\n')

# List of words not supported by spellchecker - anything with punctuation within it
with open(f"{TITLE}\\NonStandardWords.txt", 'r', encoding='utf-8') as file:
    word_list = file.readlines()
    nonStandardWords = [word.strip('\n') for word in word_list if word != '\n']
    nonStandardWordsCap = [word.capitalize() for word in nonStandardWords]

spell_dict = mySpellChecker(dict_infile, spell, replacementDict, censoredDict, nonStandardWords)
'''

def simplifyText(content):
    text_list = [str(line.text).strip() for line in content('p')]
    text_list = [l for l in text_list 
                 if not re.search(r'^Chapter|^Translator', l)
                 and l != '']
    return text_list

#re.split(pattern, string)

with open(f"{TITLE}\\Pages\\index.txt", 'r', encoding='utf-8') as infile:
    soup = BeautifulSoup(infile.read(), features="lxml")
    
if 'readlightnovel' in URL_BASE:
    contents
    contents_class = "tab-content"
elif 'wuxiaworld.online' in URL_BASE:
    contents_class = "chapter-list"
elif 'wuxiaworld.site' in URL_BASE:
    contents_class = "page-content-listing single-page"
else:
    raise ValueError("Unknown website. Please add class of chapter contents.")

chapterlist = soup.find("div", class_=contents_class)

alist = chapterlist("a", limit=CHAPTER_LIMIT)

if 'wuxiaworld' in URL_BASE:
    alist.reverse()
    
chapter_dict = {}
chapter_template = {
    'title': None,
    'url': None,
    'soup_object': None,
    'text': None
    }

# Get list of chapter pages from index
for chapter in alist:
    link = chapter.get('href')
    tmp = chapter.string.strip(' \n\t').split(' ')
    i, chapter_name = int(tmp[1].strip(':')), ' '.join(tmp[2:])
    
    try:
        tmp = chapter_name.index('Only')
    except ValueError:
        pass
    else:
        chapter_name = chapter_name[:tmp]
    
    chapter_name = chapter_name.strip('1234567890/ ()-')
    
    if chapter_name == '':
        chapter_name = None
    
    chapter_dict[i] = chapter_template.copy()
    chapter_dict[i]['title'] = chapter_name
    chapter_dict[i]['url'] = link
    print(i, chapter_name, link)

chapter_count = len(chapter_dict)

# Extract text from chapter pages
i = 0 if CHAPTER_LIMIT == 0 else chapter_count - CHAPTER_LIMIT
for chapter in chapter_dict.values():
    i += 1
    invalid_symbols = "\\/?%*:|\"<>."  
    fix_fn = lambda s: 'TBD' if chapter['title'] == None else s.translate({ord(j): None for j in invalid_symbols})
    chapter['title'] = fix_fn(chapter['title'])
    with open(f"{TITLE}\\Pages\\{i}. {chapter['title']}.txt", 
              'r', encoding='utf-8') as infile:
        chapter['soup_object'] = BeautifulSoup(infile.read(), features="lxml")
    text_area = chapter['soup_object'].find(class_="content-area")
    
    print(f'\n{i}, {chapter["title"]}')
    #chapter['title'] = spell_dict.fixCommonErrors([chapter['title'])[0]
    #chapter['text'] = spell_dict.fixCommonErrors(simplifyText(text_area))
    chapter['text'] = simplifyText(text_area)
    if chapter['text'] == '':
        raise ValueError(f'Text class incorrectly configured. No text returned for chapter {i}')

'''
# Save and close word dictionaries
with open(f"{TITLE}\\ReplacementDict.txt", 'w', encoding='utf-8') as file:
    for key, value in replacementDict.items():
        file.write(' '.join([key, value]) + '\n')
        
with open(f"{TITLE}\\NonStandardWords.txt", 'w', encoding='utf-8') as file:
    for word in sorted(list(set(nonStandardWords))):
        file.write(word + '\n')
        
dict_infile.close()
with open(filename, 'r', encoding='utf-8') as infile:
    words = infile.readlines()
    
with open(filename, 'w', encoding='utf-8') as infile:
    for word in sorted(words):
        infile.write(word)
'''
# Compact chapters
chapter_dict_tmp = {}
dummy_chapter = chapter_template.copy()
dummy_chapter['title'] = ''
i = 1
for chapter in chapter_dict.values():
    name_curr = chapter['title']
    
    name_prev = chapter_dict_tmp.get(i-1, dummy_chapter)['title']
    
    # Remove 'The ' from beginning of chapter titles for consistancy
    if len(name_prev) >= 4 and name_prev[:4] == 'The ':
        name_prev = name_prev[4:]
    if len(name_curr) >= 4 and name_curr[:4] == 'The ':
        name_curr = name_curr[4:]

    # If the chapter title is a repeat, append text to previous chapter
    if name_prev.strip('!') == name_curr.strip('!') and name_curr != 'TBD':
        chapter_dict_tmp[i-1]['text'] += chapter['text']
    else:
        chapter_dict_tmp[i] = chapter
        i += 1
        
chapter_dict = chapter_dict_tmp
    
"""----- Create Document -----"""
document = docx.Document('StyleDoc.docx')


#Title Page
document.add_paragraph('\n' * 5)
document.add_heading(TITLE, 0)
document.add_page_break()

#Contents Page
document.add_heading('Contents', 1)

for key, value in chapter_dict.items():
    document.add_paragraph(f"{key}. {value['title']}")
    
document.add_page_break()

# Write all chapters
for key, value in chapter_dict.items():
    document.add_heading(f'Chapter {key}', level=2)
    #document.add_heading(f"{key}. {value['title']}", level=2)
    
    for line in value['text']:
        document.add_paragraph(line)
    
    document.add_page_break()

document.save(f"{TITLE}\\{TITLE}.docx")
print('Document successfully created!')